import TweetQuery as tq
import pandas as pd
import requests
import matplotlib.pyplot as plt
from wordcloud import WordCloud, STOPWORDS, ImageColorGenerator
import re
from PIL import Image
import numpy as np
from docx import Document
from docx.shared import Inches
document = Document()
document.add_heading("Auto Generated Report")



tweet_df = pd.DataFrame(tq.Tweetquery('rash'))
# print(tweet_df.head())
# tweet_list = (tq.Tweetquery('Rash'))

text = " ".join(review for review in tweet_df.text)

text = re.sub(r'#\S+', '', text, flags=re.MULTILINE)
text = re.sub(r'http\S+', '', text, flags=re.MULTILINE)
text = re.sub(r'.RT\S+', '', text, flags=re.MULTILINE)
text = re.sub(r'@\S+', '', text, flags=re.MULTILINE)
print(text)
stopwords = set(STOPWORDS)
stopwords.update(["drink", "now", "wine", "flavor", "flavors","RT","Rashmika","Mandanna","cruz","know"])

# mask = np.array(Image.open(requests.get('http://www.clker.com/cliparts/O/i/x/Y/q/P/yellow-house-hi.png', stream=True).raw))
mask = np.array(Image.open('rash1.png'))

wordcloud = WordCloud(stopwords=stopwords, background_color="white",max_words=100,mask=mask).generate(text)

image_colors = ImageColorGenerator(mask)
plt.figure(figsize=[7,7])
plt.imshow(wordcloud.recolor(color_func=image_colors), interpolation="bilinear")
plt.axis("off")

plt.title('Summarization of Rashmika Tweets.')
plt.savefig('Rush.png')
document.add_picture('Rush.png', width=Inches(5))
document.save("Report")
wordcloud.to_file("Rush.png")
plt.imshow(wordcloud, interpolation='bilinear')
plt.axis("off")
api.update_with_media(filename='Rush.png',status='Rashmika Tweet summary')
plt.show()